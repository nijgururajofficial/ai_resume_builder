import logging
import os
import tempfile
import math
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx2pdf import convert

class PdfDocxGenerator:
    """
    Converts a specific Markdown structure into a visually styled, professional DOCX,
    and then converts that DOCX into a high-fidelity PDF.
    """
    def __init__(self, markdown_content: str):
        self.markdown = markdown_content
        self.font_name = 'Calibri'

    def _set_paragraph_border(self, paragraph):
        p_pr = paragraph._p.get_or_add_pPr()
        p_bdr = OxmlElement('w:pBdr')
        p_pr.append(p_bdr)
        bottom_bdr = OxmlElement('w:bottom')
        bottom_bdr.set(qn('w:val'), 'single')
        bottom_bdr.set(qn('w:sz'), '4')
        bottom_bdr.set(qn('w:space'), '1')
        bottom_bdr.set(qn('w:color'), '000000')
        p_bdr.append(bottom_bdr)

    def _add_skills_section(self, doc, lines, index):
        """Adds the SKILLS section as a two-column table."""
        # 1. Add the "SKILLS" header with a border
        p = doc.add_paragraph()
        run = p.add_run("SKILLS")
        run.font.bold = True
        run.font.size = Pt(11)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        self._set_paragraph_border(p)
        
        # 2. Collect all skill lines immediately following the header
        skill_lines = []
        i = index + 1
        while i < len(lines) and lines[i].strip().startswith('**'):
            skill_lines.append(lines[i].strip().replace('**', ''))
            i += 1
        
        if not skill_lines:
            return i

        # 3. Create a two-column table
        num_skills = len(skill_lines)
        num_rows = math.ceil(num_skills / 2)
        table = doc.add_table(rows=num_rows, cols=2)
        table.autofit = False
        table.columns[0].width = Inches(3.75)
        table.columns[1].width = Inches(3.75)

        # 4. Populate the table
        for row in range(num_rows):
            # Populate left column
            left_index = row
            raw_category, raw_skills = skill_lines[left_index].split(':', 1)
            p_left = table.cell(row, 0).paragraphs[0]
            p_left.paragraph_format.space_after = Pt(2)
            p_left.add_run(f"{raw_category.replace('**', '').strip()}:").bold = True
            p_left.add_run(f" {raw_skills.strip()}")

            # Populate right column if it exists
            right_index = row + num_rows
            if right_index < num_skills:
                raw_category, raw_skills = skill_lines[right_index].split(':', 1)
                p_right = table.cell(row, 1).paragraphs[0]
                p_right.paragraph_format.space_after = Pt(2)
                p_right.add_run(f"{raw_category.replace('**', '').strip()}:").bold = True
                p_right.add_run(f" {raw_skills.strip()}")
        
        return i

    def to_docx(self, output_path: str):
        logging.info(f"Generating styled DOCX file at: {output_path}")
        doc = Document()
        for section in doc.sections:
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
            section.top_margin = Inches(0.3)
            section.bottom_margin = Inches(0.3)

        doc.styles['Normal'].font.name = self.font_name
        doc.styles['Normal'].font.size = Pt(10.5)

        lines = self.markdown.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if not line:
                i += 1
                continue

            if line.startswith('# '):
                p = doc.add_paragraph()
                run = p.add_run(line.replace('# ', ''))
                run.font.size = Pt(22)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(2)
            
            elif ' | ' in line and '@' in line:
                p = doc.add_paragraph(line)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(8)
                self._set_paragraph_border(p)

            elif line.startswith('## '):
                # --- FIX APPLIED HERE: Simplified SKILLS section logic ---
                if "SKILLS" in line.upper():
                    # 1. Add the "SKILLS" header with a border
                    i = self._add_skills_section(doc, lines, i)
                    continue # Skip to the next major section
                else:
                    # Logic for other headers (Experience, Projects, etc.)
                    p = doc.add_paragraph()
                    run = p.add_run(line.replace('## ', '').upper())
                    run.font.bold = True
                    run.font.size = Pt(11)
                    p.paragraph_format.space_before = Pt(5)
                    p.paragraph_format.space_after = Pt(1)
                    self._set_paragraph_border(p)

            # Logic for two-column Experience headers
            elif '|||' in line:
                table = doc.add_table(rows=1, cols=2)
                table.autofit = False
                table.columns[0].width = Inches(5.0)
                table.columns[1].width = Inches(2.5)
                left_content, right_content = [p.strip() for p in line.split('|||')]
                left_cell, right_cell = table.rows[0].cells
                left_p = left_cell.paragraphs[0]
                cleaned_left = left_content.replace('**', '').strip()
                parts = cleaned_left.split('|')
                left_p.add_run(parts[0].strip()).bold = True
                if len(parts) > 1:
                    left_p.add_run(f" | {parts[1].strip()}")
                right_p = right_cell.paragraphs[0]
                right_p.text = right_content.strip()
                right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for cell in [left_cell, right_cell]:
                    cell.paragraphs[0].paragraph_format.space_before = Pt(0)
                    cell.paragraphs[0].paragraph_format.space_after = Pt(0)

            # Logic for single-line Project and Education headers
            elif line.startswith('**') and '|' in line:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                parts = line.split('|')
                p.add_run(parts[0].replace('**', '').strip()).bold = True
                if len(parts) > 1:
                    p.add_run(' | ')
                    p.add_run(parts[1].replace('*', '').strip()).italic = True

            elif line.startswith('- '):
                p = doc.add_paragraph(style='List Bullet')
                p.text = line[2:]
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
            
            # Default handler for simple lines (like the Institution name)
            else:
                p = doc.add_paragraph()
                if line.startswith('*') and line.endswith('*'):
                    p.add_run(line.replace('*', '')).italic = True
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(4)
                else:
                    p.add_paragraph(line)

            i += 1
        
        doc.save(output_path)
        logging.info("Styled DOCX generation complete.")

    # --- RE-ENGINEERED PDF METHOD ---
    # def to_pdf(self, output_path: str):
    #     """Creates a visually styled .pdf file using xhtml2pdf and advanced CSS."""
    #     logging.info(f"Generating styled PDF file at: {output_path}")

    #     html_lines = []
    #     lines = self.markdown.split('\n')
    #     i = 0
    #     while i < len(lines):
    #         line = lines[i].strip()
    #         if not line:
    #             i += 1
    #             continue

    #         if line.startswith('# '):
    #             html_lines.append(f"<h1>{line.replace('# ', '')}</h1>")
    #         elif ' | ' in line and '@' in line:
    #             html_lines.append(f"<p class='contact'>{line}</p>")
    #         elif line.startswith('## '):
    #             # Special handling for the skills section to create a grid
    #             if "SKILLS" in line.upper():
    #                 html_lines.append(f"<h2>SKILLS</h2>")
    #                 html_lines.append('<div class="skills-grid">')
    #                 i += 1
    #                 # Loop through all subsequent skill lines
    #                 while i < len(lines) and lines[i].strip().startswith('**'):
    #                     category, skill_list = lines[i].strip().split(':', 1)
    #                     html_lines.append(f"<p><b>{category.replace('**', '')}:</b>{skill_list}</p>")
    #                     i += 1
    #                 html_lines.append('</div>')
    #                 continue # Skip the increment at the end of the main loop
    #             else:
    #                 html_lines.append(f"<h2>{line.replace('## ', '').upper()}</h2>")

    #         # Handles the single-line header format: "Left Part ||| Right Part"
    #         elif line.startswith('**') and '|||' in line:
    #             left_content, right_content = [p.strip() for p in line.split('|||')]
    #             left_content = left_content.replace('**', '')
    #             left_title, left_company = (left_content.split('|', 1) + [''])[:2]
                
    #             html_lines.append('<div class="item-header">')
    #             html_lines.append(f'  <div class="item-left"><b>{left_title.strip()}</b>' + (f' | {left_company.strip()}' if left_company.strip() else '') + '</div>')
    #             if right_content.startswith('*') and right_content.endswith('*'):
    #                 html_lines.append(f'  <div class="item-right"><em>{right_content[1:-1]}</em></div>')
    #             else:
    #                 html_lines.append(f'  <div class="item-right">{right_content}</div>')
    #             html_lines.append('</div>')

    #         # Handles bullet points with proper <ul> list wrapping
    #         elif line.startswith('- '):
    #             # Start the list if the previous line was not a list item
    #             if not (i > 0 and lines[i-1].strip().startswith('- ')):
    #                 html_lines.append('<ul>')
                
    #             html_lines.append(f"<li>{line[2:]}</li>")
                
    #             # End the list if the next line is not a list item
    #             if not (i + 1 < len(lines) and lines[i+1].strip().startswith('- ')):
    #                 html_lines.append('</ul>')
            
    #         i += 1

    #     html_content = "\n".join(html_lines)

    #     # CSS is now specifically designed for the target layout
    #     css_string = f"""
    #         @page {{ 
    #             margin: 0.5in; 
    #         }}
    #         body {{ 
    #             font-family: '{self.font_name}', sans-serif; 
    #             font-size: 10.5pt; 
    #             color: #000000; 
    #         }}
    #         h1 {{ 
    #             font-size: 22pt; 
    #             text-align: center; 
    #             margin: 0; 
    #             font-weight: normal; 
    #         }}
    #         .contact {{ 
    #             text-align: center; 
    #             margin-bottom: 6pt; 
    #             font-size: 10pt; 
    #         }}
    #         h2 {{ 
    #             font-size: 11pt; 
    #             font-weight: bold;
    #             margin: 8pt 0 4pt 0; 
    #             padding-bottom: 2px;
    #             border-bottom: 0.5pt solid #000; 
    #         }}
    #         .skills-grid {{
    #             display: grid;
    #             grid-template-columns: 1fr 1fr;
    #             gap: 0px 15px; /* Row gap 0, Column gap 15px */
    #             padding-left: 5px; /* Small indent for alignment */
    #         }}
    #         .skills-grid p {{ 
    #             margin-bottom: 2pt; 
    #         }}
    #         .item-header {{
    #             display: flex;
    #             justify-content: space-between;
    #             width: 100%;
    #             margin-top: 4pt;
    #             margin-bottom: 2pt;
    #         }}
    #         .item-left {{ text-align: left; }}
    #         .item-right {{ text-align: right; }}
    #         em {{ font-style: italic; }}
    #         ul {{ 
    #             margin: 0; 
    #             padding-left: 20px; 
    #             list-style-type: disc;
    #         }}
    #         li {{ 
    #             margin-bottom: 2pt; 
    #         }}
    #         p {{ 
    #             margin: 0; 
    #             padding: 0; 
    #         }}
    #     """

    #     final_html = f"<html><head><style>{css_string}</style></head><body>{html_content}</body></html>"

    #     with open(output_path, "w+b") as pdf_file:
    #         pisa_status = pisa.CreatePDF(src=final_html, dest=pdf_file)
        
    #     if pisa_status.err:
    #         logging.error(f"Error generating PDF: {pisa_status.err}")
    #     else:
    #         logging.info("Styled PDF generation complete.")

    def to_pdf(self, output_path: str):
        """
        Creates a high-fidelity PDF by first generating a DOCX and then converting it.
        This method replaces the previous HTML-based approach.
        """
        logging.info(f"Generating high-fidelity PDF at: {output_path}")
        
        # Use a temporary directory to avoid leaving behind intermediate files
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_docx_path = os.path.join(temp_dir, "temp_resume.docx")
            
            try:
                # Step 1: Generate the perfect DOCX file using the unchanged method above.
                self.to_docx(temp_docx_path)
                
                # Step 2: Convert that DOCX file to PDF.
                convert(temp_docx_path, output_path)
                
                logging.info("High-fidelity PDF generation complete.")
            
            except Exception as e:
                logging.error(f"An error occurred during PDF conversion: {e}")
                logging.error(
                    "Please ensure Microsoft Word (on Windows) or LibreOffice (on macOS/Linux) is installed."
                )