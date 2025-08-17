import os
import json
import re
import textwrap
from typing import TypedDict, List
from dotenv import load_dotenv
import google.generativeai as genai
from langgraph.graph import StateGraph, START, END
from langchain_core.messages import HumanMessage
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from reportlab.lib.pagesizes import LETTER
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

# ==============================
# 1. Load API key
# ==============================
load_dotenv()
os.environ["GOOGLE_API_KEY"] = "AIzaSyDFmAT7k6pVujNqyAubBKO9NFCj7ZpMJYY"
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

# ==============================
# 2. JSON parser
# ==============================
def parse_ai_json(ai_text):
    """Extract JSON from AI output."""
    json_match = re.search(r"```json\s*(\{.*?\})\s*```", ai_text, re.DOTALL)
    json_str = json_match.group(1) if json_match else ai_text.strip()
    try:
        return json.loads(json_str)
    except json.JSONDecodeError as e:
        print(f"⚠️ JSON parsing error: {e}")
        return None

# ==============================
# 3. Agent state
# ==============================
class AgentState(TypedDict):
    messages: List[HumanMessage]
    job_description: str
    resume_data: dict
    ats_keywords: dict
    keyword_gaps: dict
    keyword_suggestions: dict
    resume_format_md: str

model = genai.GenerativeModel("models/gemini-2.5-pro")

# ==============================
# 4. Load files
# ==============================
def load_files(state: AgentState) -> AgentState:
    try:
        with open("summary.json", "r", encoding="utf-8") as f:
            job_description = f.read()
        with open("parsed_resume.json", "r", encoding="utf-8") as f:
            resume_data = json.load(f)
        with open("resume_format.md", "r", encoding="utf-8") as f:
            resume_format_md = f.read()
    except Exception as e:
        print(f"❌ Error reading files: {e}")
        return state

    print("✅ Job description, resume, and format markdown loaded successfully.")
    state["job_description"] = job_description
    state["resume_data"] = resume_data
    state["resume_format_md"] = resume_format_md
    return state

# ==============================
# 5. Extract ATS keywords
# ==============================
def extract_keywords(state: AgentState) -> AgentState:
    job_desc = state.get("job_description", "")

    ats_prompt = f"""
You are an expert in Applicant Tracking Systems (ATS) and resume optimization.

Identify the MOST important keywords and phrases in the job description that will help a resume pass ATS filters.
For each keyword, include:
- exact term
- frequency
- category: Technical Skill | Tool/Software | Soft Skill | Qualification | Other

Return ONLY valid JSON:

{{
  "keywords": [
    {{
      "term": "example term",
      "frequency": 1,
      "category": "Technical Skill"
    }}
  ]
}}

Job Description:
{job_desc}
"""

    print("🤖 Extracting ATS keywords...")
    response = model.generate_content(ats_prompt)
    ai_text = response.text if response else "{}"

    ats_keywords_json = parse_ai_json(ai_text)
    if ats_keywords_json:
        state["ats_keywords"] = ats_keywords_json
        with open("ats_keywords.json", "w", encoding="utf-8") as f:
            json.dump(ats_keywords_json, f, indent=2, ensure_ascii=False)
        print("✅ ATS keywords saved to 'ats_keywords.json'")
    else:
        print("⚠️ Failed to parse ATS keywords JSON.")
    return state

# ==============================
# 6. Compare Resume vs Keywords & Suggest Insertions
# ==============================
def compare_resume_to_keywords(state: AgentState) -> AgentState:
    ats_keywords = state.get("ats_keywords", {}).get("keywords", [])
    resume_data = state.get("resume_data", {})
    resume_text = json.dumps(resume_data, ensure_ascii=False).lower()

    missing_keywords = []
    for kw in ats_keywords:
        term = kw.get("term", "").lower()
        if resume_text.count(term) < kw.get("frequency", 1):
            missing_keywords.append({
                "term": kw.get("term"),
                "category": kw.get("category"),
                "required_frequency": kw.get("frequency"),
                "resume_count": resume_text.count(term)
            })

    state["keyword_gaps"] = {"missing_or_underused_keywords": missing_keywords}

    suggestion_prompt = f"""
You are an ATS optimization expert.

Given the resume JSON and missing keywords, suggest where to insert each keyword naturally
without breaking grammar or facts. Return JSON:

{{
  "suggestions": [
    {{
      "term": "keyword",
      "section": "Resume Section",
      "suggestion": "Example sentence or bullet with keyword"
    }}
  ]
}}

Resume JSON:
{json.dumps(resume_data, ensure_ascii=False)}

Missing Keywords:
{json.dumps(missing_keywords, ensure_ascii=False)}
"""

    print("🤖 Generating insertion suggestions...")
    response = model.generate_content(suggestion_prompt)
    ai_text = response.text if response else "{}"

    suggestions_json = parse_ai_json(ai_text)
    if suggestions_json:
        state["keyword_suggestions"] = suggestions_json
        with open("resume_keyword_suggestions.json", "w", encoding="utf-8") as f:
            json.dump(suggestions_json, f, indent=2, ensure_ascii=False)
        print("✅ Keyword suggestions saved to 'resume_keyword_suggestions.json'")
    else:
        print("⚠️ Failed to parse keyword suggestions JSON.")

    return state

# ==============================
# 7. Generate Filled DOCX & PDF
# ==============================
def generate_resume_files(state: AgentState) -> AgentState:
    filled_md = state.get("resume_format_md", "")
    resume_data = state.get("resume_data", {})
    suggestions = state.get("keyword_suggestions", {}).get("suggestions", [])

    # Replace placeholders in markdown with actual resume data
    for key, value in resume_data.items():
        if isinstance(value, list):
            value_str = "\n".join([f"- {v}" for v in value])
        else:
            value_str = str(value)
        filled_md = filled_md.replace(f"{{{{ {key} }}}}", value_str)

    # --- DOCX ---
    doc = Document()
    for line in filled_md.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("# "):
            p = doc.add_paragraph()
            run = p.add_run(line[2:])
            run.bold = True
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif line.startswith("## "):
            p = doc.add_paragraph()
            run = p.add_run(line[3:])
            run.bold = True
        elif line.startswith("### "):
            p = doc.add_paragraph()
            run = p.add_run(line[4:])
            run.bold = True
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)

    # ATS keyword suggestions
    doc.add_page_break()
    doc.add_paragraph("ATS Keyword Suggestions", style="Heading 1")
    for s in suggestions:
        doc.add_paragraph(f"Term: {s.get('term','')}")
        doc.add_paragraph(f"Section: {s.get('section','')}")
        doc.add_paragraph(f"Suggestion: {s.get('suggestion','')}")
        doc.add_paragraph("")

    doc.save("final_resume.docx")

    # --- PDF ---
    pdf = SimpleDocTemplate("final_resume.pdf", pagesize=LETTER)
    styles = getSampleStyleSheet()
    normal_style = styles["Normal"]
    heading_style = ParagraphStyle('Heading1', parent=styles['Heading1'], spaceAfter=12)
    elements = []

    for line in filled_md.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("# "):
            elements.append(Paragraph(line[2:], heading_style))
        elif line.startswith("## "):
            elements.append(Paragraph(line[3:], heading_style))
        elif line.startswith("### "):
            elements.append(Paragraph(line[4:], heading_style))
        elif line.startswith("- "):
            elements.append(Paragraph(f"• {line[2:]}", normal_style))
        else:
            elements.append(Paragraph(line, normal_style))
        elements.append(Spacer(1, 4))

    elements.append(Spacer(1, 12))
    elements.append(Paragraph("ATS Keyword Suggestions", heading_style))
    for s in suggestions:
        elements.append(Paragraph(f"<b>Term:</b> {s.get('term','')}", normal_style))
        elements.append(Paragraph(f"<b>Section:</b> {s.get('section','')}", normal_style))
        elements.append(Paragraph(f"<b>Suggestion:</b> {s.get('suggestion','')}", normal_style))
        elements.append(Spacer(1, 6))

    pdf.build(elements)

    print("✅ final_resume.docx and final_resume.pdf generated.")
    return state

# ==============================
# 8. Build Agent Graph
# ==============================
graph = StateGraph(AgentState)
graph.add_node("load_files", load_files)
graph.add_node("extract_keywords", extract_keywords)
graph.add_node("compare_resume_to_keywords", compare_resume_to_keywords)
graph.add_node("generate_resume_files", generate_resume_files)
graph.add_edge(START, "load_files")
graph.add_edge("load_files", "extract_keywords")
graph.add_edge("extract_keywords", "compare_resume_to_keywords")
graph.add_edge("compare_resume_to_keywords", "generate_resume_files")
graph.add_edge("generate_resume_files", END)
agent = graph.compile()

# ==============================
# 9. Run Agent
# ==============================
if __name__ == "__main__":
    agent.invoke({"messages": []})